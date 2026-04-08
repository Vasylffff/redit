"""Build Assessment 2 Word document from ASSESSMENT_2_REPORT.md with embedded figures."""
import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

REPORT_PATH = "ASSESSMENT_2_REPORT.md"
FIGURES_DIR = "data/analysis/reddit/figures"
OUTPUT_PATH = "Assessment_2_Final.docx"

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# Styles
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles["Heading %d" % level]
    hs.font.name = "Calibri"
    hs.font.color.rgb = RGBColor(0, 0, 0)

# Read markdown
with open(REPORT_PATH, "r", encoding="utf-8") as f:
    lines = f.readlines()

# Figure mapping
figure_files = {
    "fig1": "fig1_topic_trajectories.png",
    "fig2": "fig2_roc_emergence_models.png",
    "fig3": "fig3_feature_importance.png",
    "fig4": "fig4_depth_vs_roc.png",
    "fig5": "fig5_pipeline_summary.png",
    "fig6": "fig6_revival_rates.png",
    "fig7": "fig7_model_heatmap.png",
    "fig8": "fig8_subreddit_spread.png",
    "fig9": "fig9_confusion_matrix.png",
    "fig10": "fig10_gbm_learning_rate.png",
    "fig11": "fig11_data_coverage.png",
    "fig12": "fig12_ongoing_vs_oneshot.png",
    "fig13": "fig13_transition_matrix.png",
    "fig14": "fig14_target_vs_roc.png",
}

figure_captions = {
    "Figure 1": "Topic lifecycle trajectories showing ongoing stories vs one-shot events across 13 days of data collection.",
    "Figure 2": "ROC curves comparing five classifiers on topic emergence detection (1-3 posts predicting growth to 5+).",
    "Figure 3": "Feature importance for topic emergence detection using Random Forest. Post count dominates at 32.8%.",
    "Figure 4": "Overfitting analysis: deeper trees consistently perform worse. Random Forest peaks at depth=4, Decision Tree at depth=3.",
    "Figure 5": "Complete topic lifecycle prediction pipeline showing ROC AUC for each task. Green = excellent (>0.9), orange = good (>0.8), red = moderate.",
    "Figure 6": "Left: false death rates decrease with stricter definitions. Right: bigger topics have higher revival rates (44.8% for 8-11 peak posts).",
    "Figure 7": "Model performance heatmap across all tasks. Bold values indicate best model per task. No single model dominates.",
    "Figure 8": "Left: r/politics breaks stories first most often. Right: top cross-subreddit spread routes.",
    "Figure 9": "Confusion matrix for emergence detection at 10% probability threshold.",
    "Figure 10": "Gradient Boosting performance peaks at learning rate 0.02. Higher rates cause overfitting.",
    "Figure 11": "Daily data collection coverage by subreddit showing the 13-day collection period.",
    "Figure 12": "Feature ratios distinguishing ongoing stories from one-shot events. Multiple peaks is the strongest signal at 15.5x.",
    "Figure 13": "Topic state transition matrix showing daily state changes. Dead topics stay dead 94% of the time.",
    "Figure 14": "Emergence detection ROC AUC at different growth targets (3+ to 15+ posts).",
}


def add_figure(fig_key, caption_key):
    """Add a figure with caption if the file exists."""
    for fk, fname in figure_files.items():
        if fk in fig_key.lower().replace(" ", "").replace("figure", "fig"):
            fpath = os.path.join(FIGURES_DIR, fname)
            if os.path.exists(fpath):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(fpath, width=Inches(5.5))
                # Caption
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.paragraph_format.space_after = Pt(12)
                run = cap.add_run(caption_key)
                run.font.size = Pt(9)
                run.font.italic = True
                return True
    return False


def add_table(headers, rows_data):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
    # Data
    for ri, row in enumerate(rows_data):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()  # spacing


# ================================================================
# BUILD DOCUMENT
# ================================================================

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Predicting Reddit Topic Lifecycles:\nEmergence, Spread, and Death")
run.font.size = Pt(18)
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Exploring AI: Understanding and Applications (SPC4004)\nAssessment 2 -- Code Generation Project")
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("GitHub: https://github.com/Vasylffff/redit")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0, 102, 204)

doc.add_page_break()

# ================================================================
# SECTION 1: Problem Definition
# ================================================================
doc.add_heading("1. Problem Definition & Dataset Justification", level=1)

doc.add_paragraph(
    "This project builds a complete topic lifecycle prediction system for Reddit. "
    "Rather than predicting individual posts, it tracks how topics -- represented as "
    "co-occurring word pairs in post titles -- emerge, spread across communities, and die. "
    "The system is entirely content-agnostic: the same algorithm detects a political scandal, "
    "a game launch, or a viral meme using only engagement signals, never reading the actual content."
)

doc.add_paragraph(
    "The machine learning task combines multiple problems: binary classification (will a topic grow?), "
    "multi-class classification (what state will it be in tomorrow?), and regression (how many days "
    "until death?). Different algorithms proved appropriate for different lifecycle stages, and a key "
    "finding was that the simplest model -- Logistic Regression -- outperforms complex ensembles on "
    "the core detection task."
)

doc.add_heading("Dataset", level=2)

doc.add_paragraph(
    "Entirely self-collected over 13 days (26 March -- 7 April 2026) using Reddit's public JSON "
    "endpoints across five subreddits: r/technology, r/news, r/worldnews, r/politics, and r/Games. "
    "Data was collected hourly via Windows Task Scheduler from two machines, then merged."
)

add_table(
    ["Metric", "Value"],
    [
        ["Post snapshots", "216,944"],
        ["Comment snapshots", "972,353"],
        ["Unique posts tracked", "6,826"],
        ["Collection frequency", "Hourly"],
        ["Raw JSON files", "5,140"],
        ["Co-occurrence pairs analysed", "200,000+ per day"],
    ]
)

doc.add_paragraph(
    "Posts are tracked through lifecycle states -- surging, alive, cooling, dying, dead -- "
    "computed from upvote and comment velocity thresholds per subreddit. These serve as natural "
    "labels without manual annotation. Topics are represented as two-word pairs extracted from "
    "titles (e.g. \"russian+tanker\", \"crimson+desert\"), which capture specific stories rather "
    "than generic vocabulary (Szabo and Huberman, 2010)."
)

add_figure("fig11", "Figure 11: Daily data collection coverage by subreddit showing the 13-day collection period.")

# ================================================================
# SECTION 2: Initial Code
# ================================================================
doc.add_heading("2. Initial Code & Explanation of AI Use", level=1)

doc.add_paragraph(
    "The codebase was generated using Claude Code (Anthropic, Claude Opus 4.6). Initial prompt:"
)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1.5)
run = p.add_run('"This is my project and I want you to create a folder and set this up so I could scrape Reddit from here"')
run.font.italic = True

doc.add_paragraph(
    "This produced a data collection pipeline and a five-layer Markov chain post predictor "
    "using only upvote velocity. The initial system had no topic-level prediction, no comment "
    "engagement analysis, no cross-subreddit tracking, and no sentiment features. All subsequent "
    "development used Claude Code conversationally -- the developer directed priorities, questioned "
    "results, and pushed exploration beyond what the AI initially suggested was achievable."
)

# ================================================================
# SECTION 3: Critique
# ================================================================
doc.add_heading("3. Critique of Initial Code", level=1)

doc.add_heading("Bugs", level=2)

bugs = [
    "datetime.UTC incompatible with Python 3.10 across 12 files",
    "f-string backslash syntax errors",
    "scikit-learn version incompatibility",
    "Windows Task Scheduler crashed due to WPF popups in non-interactive mode",
    "Unicode em dashes corrupted under Windows cp1252 encoding",
]
for b in bugs:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Algorithmic Limitations", level=2)

doc.add_paragraph(
    "No topic-level analysis. The initial code predicted individual posts but had no concept of "
    "\"topics\" -- it could not detect that multiple posts about the same story were related."
)

doc.add_paragraph(
    "Single-keyword approach. When topic analysis was first attempted, it used single keywords "
    "(\"trump\", \"iran\"), which are too generic. The word \"trump\" appears in 2,192 posts -- "
    "it is not a topic, it is noise. The shift to co-occurrence pairs (\"russian+tanker\", "
    "\"birthright+citizenship\") was a critical improvement."
)

doc.add_paragraph(
    "Random Forest everywhere. The initial code defaulted to Random Forest for every task. "
    "Testing five classifiers across seven tasks revealed that Logistic Regression outperforms "
    "Random Forest on the core emergence detection task (0.860 vs 0.829 ROC AUC), indicating "
    "the signal is fundamentally linear."
)

doc.add_paragraph(
    "No death definition analysis. The initial code defined topic death as \"<2 posts for 1 day\". "
    "Analysis showed this produces a 13.1% false-death rate. A 2-consecutive-day definition "
    "reduces this to 6.8%."
)

doc.add_paragraph(
    "Default hyperparameters. All models used default settings. Tuning revealed Gradient Boosting "
    "was severely misconfigured (depth=6 caused overfitting, ROC improved from 0.713 to 0.835 at "
    "depth=2), and Decision Tree improved from 0.577 to 0.842 by reducing depth from 8 to 5."
)

add_figure("fig4", "Figure 4: Overfitting analysis -- deeper trees consistently perform worse. Random Forest peaks at depth=4, Decision Tree at depth=3.")

# ================================================================
# SECTION 4: Iterative Development
# ================================================================
doc.add_heading("4. Iterative Development & Justification", level=1)

# Phase 1
doc.add_heading("Phase 1: Making It Work", level=2)
doc.add_paragraph(
    "Iteration 1 -- Compatibility fixes. Replaced datetime.UTC with timezone.utc in 12 files. "
    "Fixed f-string syntax. Pinned scikit-learn. Fixed Task Scheduler encoding issues. "
    "Verified: system collects data hourly without crashing."
)

# Phase 2
doc.add_heading("Phase 2: Post-Level Engagement", level=2)
doc.add_paragraph(
    "Iteration 2 -- VADER sentiment. Scored 972K comments. Found negative sentiment correlates "
    "with longer post survival (alive avg -0.006 vs dead +0.045). Controversy drives engagement."
)
doc.add_paragraph(
    "Iteration 3 -- Gini coefficient. Comment upvote distribution became the strongest "
    "post-level predictor at 46% feature importance. High Gini = community consensus = survival."
)
doc.add_paragraph(
    "Iteration 4 -- Per-subreddit models. r/politics: 81% accuracy. r/Games: 64%. "
    "Each community has distinct engagement patterns."
)

# Phase 3
doc.add_heading("Phase 3: Post Prediction at Multiple Horizons", level=2)
doc.add_paragraph(
    "Iteration 5 -- Multi-horizon classifiers. Built 17 classifiers for 1h to 72h horizons. "
    "ROC AUC decays from 0.843 (1h) to 0.57 (7 days). Beyond 48 hours, accuracy paradoxically "
    "rises to 85% while ROC drops -- predicting \"everything dies\" gets high accuracy but "
    "zero discriminative power."
)
doc.add_paragraph(
    "Iteration 6 -- Surging and dead detection. Surging detection: 0.987 ROC AUC. "
    "Dead detection: 0.945."
)

# Phase 4
doc.add_heading("Phase 4: From Posts to Topics", level=2)
doc.add_paragraph(
    "Iteration 7 -- Co-occurrence pairs. Two-word pairs from the same title represent specific "
    "stories. \"birthright+citizenship\" is a Supreme Court case. \"russian+tanker\" is a naval "
    "incident. Temporal validation: 0.813 ROC AUC."
)

add_figure("fig1", "Figure 1: Topic lifecycle trajectories showing ongoing stories vs one-shot events across 13 days.")

doc.add_paragraph(
    "Iteration 8 -- Content-agnostic detection. The same algorithm catches completely different "
    "types of content: official+trailer (122 posts, game announcements), hormuz+strait (252 posts, "
    "geopolitical crisis), crimson+desert (65 posts, game launch), media+social (127 posts, tech "
    "policy debate). The model does not read content -- it tracks words appearing together and "
    "people engaging. A viral meme and a war follow identical spread patterns."
)

# Phase 5
doc.add_heading("Phase 5: Complete Topic Lifecycle", level=2)

doc.add_paragraph(
    "Iteration 9 -- \"Has it peaked?\" Given a topic's first two days of data, predict whether "
    "it has already peaked or will keep growing. ROC AUC: 0.958. The signal is simple: "
    "post_growth_d1_d2 (52.7% importance) and upvote_growth_d1_d2 (38.1%)."
)

doc.add_paragraph(
    "Iteration 10 -- Topic death prediction. \"Will this topic die tomorrow?\" ROC AUC: 0.890. "
    "Top features: current post count (46.1%), subreddit coverage (17.1%)."
)

doc.add_paragraph(
    "Iteration 11 -- Quick vs slow death. After a topic peaks, will it die in 0-1 days or "
    "survive 2+ days? ROC AUC: 0.996. The single feature: decline_rate_d1 (67.8% importance)."
)

add_figure("fig5", "Figure 5: Complete topic lifecycle prediction pipeline with ROC AUC scores. Green = excellent (>0.9), orange = good (>0.8).")

doc.add_paragraph(
    "Iteration 12 -- Subreddit spread prediction. If a topic exists in subreddit A, will it "
    "appear in subreddit B tomorrow? Per-target ROC AUC: r/politics 0.756, r/Games 0.679. "
    "Key finding: r/politics breaks stories first (47,580 times), not r/news."
)

add_figure("fig8", "Figure 8: Left: r/politics breaks stories first most often. Right: top cross-subreddit spread routes.")

doc.add_paragraph(
    "Iteration 13 -- Ongoing vs one-shot classification. When a topic drops, will it come back? "
    "ROC AUC: 0.970. The strongest signal: multiple_peaks (21.5% importance). Ongoing stories "
    "have multiple surges and dips; one-shot events peak once and die."
)

add_table(
    ["Signal", "Ongoing story", "One-shot event", "Ratio"],
    [
        ["Multiple peaks", "0.9", "0.1", "15.5x"],
        ["Consistency", "0.4", "0.1", "4.6x"],
        ["Active days", "3.2", "0.8", "4.0x"],
        ["Peak posts", "3.8", "1.6", "2.3x"],
    ]
)

add_figure("fig12", "Figure 12: Feature ratios distinguishing ongoing stories from one-shot events. Multiple peaks is the strongest signal at 15.5x.")

doc.add_paragraph(
    "Iteration 14 -- Topic death definition. Analysing 8,880 drop events revealed the 1-day "
    "death definition has 13.1% false-death rate. Bigger topics revive more often: 44.8% revival "
    "rate for topics peaking at 8-11 posts vs 22.8% for small topics. Two consecutive days "
    "reduces false deaths to 6.8%."
)

add_figure("fig6", "Figure 6: Left: false death rates decrease with stricter definitions. Right: bigger topics have higher revival rates.")

doc.add_paragraph(
    "Iteration 15 -- Noise filter. 98.6% of word pairs never grow. At 99% confidence, the "
    "model filters out 87% of all pairs as noise with 99.5% precision. ROC AUC: 0.824."
)

doc.add_paragraph(
    "Iteration 16 -- Model comparison and hyperparameter tuning. Tested five classifiers across "
    "all tasks with 36 hyperparameter configurations."
)

add_figure("fig7", "Figure 7: Model performance heatmap across all tasks. Bold = best model per task. No single model dominates.")

# ================================================================
# SECTION 5: Final Evaluation
# ================================================================
doc.add_heading("5. Final Code Evaluation and Reflection", level=1)

doc.add_heading("The Complete Lifecycle Pipeline", level=2)

add_table(
    ["Lifecycle Stage", "Task", "Best ROC AUC", "Best Model"],
    [
        ["Filter", "Discard noise (87% filtered)", "0.850", "Logistic Regression"],
        ["Birth", "Detect emerging topic (1-3 -> 5+)", "0.860", "Logistic Regression"],
        ["Growth", "Has it peaked or still growing?", "0.958", "Gradient Boosting"],
        ["Spread", "Will it reach r/politics?", "0.756", "Random Forest"],
        ["Decline", "Topic dying state detection", "0.992", "Random Forest"],
        ["Death", "Will it die tomorrow?", "0.890", "Random Forest"],
        ["Death speed", "Quick death or slow death?", "0.999", "Logistic Regression"],
        ["Revival", "Ongoing story or one-shot?", "0.970", "Random Forest"],
    ]
)

doc.add_heading("Model Comparison", level=2)

add_figure("fig2", "Figure 2: ROC curves comparing five classifiers on topic emergence detection.")

doc.add_paragraph(
    "No single model dominates. Logistic Regression wins on emergence detection (0.860 vs RF 0.829), "
    "meaning the topic emergence signal is fundamentally linear. Random Forest wins on tasks "
    "requiring complex feature interactions. Gradient Boosting is sensitive to hyperparameters."
)

doc.add_heading("Hyperparameter Tuning Impact", level=2)

add_table(
    ["Model", "Default ROC", "Tuned ROC", "Improvement"],
    [
        ["Decision Tree", "0.577", "0.842", "+0.265"],
        ["Gradient Boosting", "0.713", "0.835", "+0.122"],
        ["Random Forest", "0.820", "0.846", "+0.027"],
        ["Logistic Regression", "0.859", "0.860", "+0.001"],
    ]
)

doc.add_paragraph(
    "Overfitting is the dominant failure mode. Across all model families, shallower trees "
    "outperform deeper ones. Random Forest peaks at depth=4, Decision Tree at depth=5. "
    "Unlimited depth drops ROC to 0.698."
)

add_figure("fig10", "Figure 10: Gradient Boosting performance peaks at learning rate 0.02. Higher rates cause overfitting.")

doc.add_heading("Topic State Transitions", level=2)

add_figure("fig13", "Figure 13: Topic state transition matrix showing daily state changes. Dead topics stay dead 94% of the time.")

doc.add_paragraph(
    "Surging topics: 48% go stable next day, 30% die, only 4% stay surging. "
    "Dead topics: 94% stay dead, but 6% revive. Dying topics: 44% revive to stable -- "
    "topic death is not permanent."
)

doc.add_heading("What Doesn't Work", level=2)

doc.add_paragraph(
    "Magnitude prediction. We can detect whether a topic will grow (0.86 ROC) but cannot predict "
    "how much. Tested Szabo-Huberman log-linear model (R2=0.22), Random Forest regression (R2=0.16), "
    "power-law percentile ranges, and growth multipliers. All fail because at 1-3 posts, 99.5% of "
    "word pairs stay small."
)

doc.add_paragraph(
    "Exact timing. \"How many days until death?\" gives R2=-0.5 (worse than guessing the mean). "
    "The binary question works; regression does not."
)

doc.add_paragraph(
    "Revival timing. \"When will a dead topic come back?\" gives 0.578 ROC -- barely above random. "
    "Revival depends on external real-world events, not engagement metrics. However, predicting "
    "which topics are the type that revives works at 0.970 ROC."
)

doc.add_heading("Limitations", level=2)
limitations = [
    "13 days of data limits rare event learning (only ~62 topic growth events in test sets)",
    "Daily granularity -- hourly topic tracking would provide more signal",
    "No natural language understanding -- distinguishes ongoing from one-shot by trajectory shape only",
    "Magnitude prediction remains unsolved from early signals",
    "All findings are correlational, not causal",
]
for lim in limitations:
    doc.add_paragraph(lim, style="List Bullet")

# ================================================================
# SECTION 6: Reflection
# ================================================================
doc.add_heading("6. Reflection on AI-Assisted Coding", level=1)

doc.add_heading("Where AI Was Effective", level=2)
doc.add_paragraph(
    "Rapid prototyping of data collection, algorithm selection, and boilerplate generation. "
    "The conversational workflow enabled moving from idea to working analysis in minutes."
)

doc.add_heading("Where AI Was Wrong or Misleading", level=2)

doc.add_paragraph(
    "Default model assumptions. AI defaulted to Random Forest for every task. Testing revealed "
    "Logistic Regression outperforms it on the core task -- the AI never suggested trying "
    "simpler models first."
)

doc.add_paragraph(
    "Activity state labels for topics. AI suggested using post-level activity states as features "
    "for topic prediction. These carry zero signal at topic level (alive_ratio showed 0.9x ratio "
    "between growing and non-growing topics). The developer identified this as redundant."
)

doc.add_paragraph(
    "Premature ceiling claims. AI declared \"we've hit the ceiling\" multiple times. The developer "
    "pushed past each one, discovering the topic lifecycle pipeline, subreddit spread prediction, "
    "and the ongoing-vs-one-shot classification."
)

doc.add_paragraph(
    "Hyperparameter negligence. AI used near-default parameters throughout. Tuning revealed "
    "Decision Tree improved by +0.265 and Gradient Boosting by +0.122."
)

doc.add_heading("Validation Approach", level=2)
validations = [
    "Temporal validation (trained on past, tested on future) rather than cross-validation",
    "Feature importance analysis to verify models learn real patterns",
    "Confusion matrices to understand failure modes",
    "Multiple model comparison to avoid algorithm bias",
    "Testing established research models (Szabo-Huberman) against our data",
    "Honest reporting of what does not work alongside what does",
]
for v in validations:
    doc.add_paragraph(v, style="List Bullet")

doc.add_heading("Ethical Considerations", level=2)
ethics = [
    "1.1-second rate limiting between API requests",
    "Aggregate analysis only, no individual user profiling",
    "Public data only (Reddit JSON endpoints, no authentication required)",
    "Transparent documentation of AI tool usage throughout",
]
for e in ethics:
    doc.add_paragraph(e, style="List Bullet")

# ================================================================
# References
# ================================================================
doc.add_heading("References", level=1)

refs = [
    "Hutto, C.J. and Gilbert, E.E. (2014) 'VADER: A Parsimonious Rule-based Model for Sentiment Analysis of Social Media Text', Proceedings of the International AAAI Conference on Web and Social Media.",
    "Pedregosa, F. et al. (2011) 'Scikit-learn: Machine Learning in Python', Journal of Machine Learning Research, 12, pp. 2825-2830.",
    "Szabo, G. and Huberman, B.A. (2010) 'Predicting the Popularity of Online Content', Communications of the ACM, 53(8), pp. 80-88.",
]
for r in refs:
    p = doc.add_paragraph(r)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.left_indent = Cm(1.27)

# Save
doc.save(OUTPUT_PATH)
print("Saved: %s" % OUTPUT_PATH)
print("Done!")
