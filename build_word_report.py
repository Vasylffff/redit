"""Build Word document for Assessment 2 submission."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Predicting Reddit Post Survival Using\nComment Engagement and Sentiment Analysis')
run.bold = True
run.font.size = Pt(16)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Exploring AI: Understanding and Applications (SPC4004)\nAssessment 2 \u2013 Code Generation Project')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('GitHub Repository: ').bold = True
p.add_run('https://github.com/Vasylffff/redit')

# ===== SECTION 1 =====
doc.add_heading('1. Problem Definition & Dataset Justification', level=2)

doc.add_paragraph(
    'This project addresses a binary classification problem: given a Reddit post\'s current metrics, '
    'can we predict whether it will remain active or decay within a specified time horizon?'
)

doc.add_paragraph(
    'Reddit posts follow a measurable lifecycle, transitioning through states: emerging, surging, alive, '
    'cooling, dying, and dead. These states are determined by upvote and comment velocity thresholds '
    'computed empirically per subreddit. Predicting this trajectory has applications in content strategy, '
    'news monitoring, and understanding online engagement dynamics.'
)

doc.add_paragraph(
    'The dataset was self-collected using Reddit\'s public JSON endpoints, requiring no API authentication. '
    'Over 12 days (26 March to 6 April 2026), the system collected hourly snapshots from five subreddits: '
    'r/technology, r/news, r/worldnews, r/politics, and r/Games. The final dataset comprises 185,048 post '
    'snapshots, 677,569 comment snapshots, and 6,292 unique post lifecycles. Each snapshot records upvote '
    'count, comment count, velocity metrics, rank position, and computed activity state. Comment snapshots '
    'include full text, enabling natural language processing.'
)

doc.add_paragraph(
    'This dataset was chosen because it provides naturally labelled training data: lifecycle states serve '
    'as ground truth without requiring manual annotation. The time-series structure enables both '
    'single-point classification and trajectory prediction across multiple horizons.'
)

# ===== SECTION 2 =====
doc.add_heading('2. Initial Code & Explanation of AI Use', level=2)

doc.add_paragraph(
    'The initial codebase was generated using Claude Code (Anthropic\'s CLI, powered by Claude Opus 4.6). '
    'The primary prompt was:'
)

quote = doc.add_paragraph(
    '"This is my project and I want you to create a folder and set this up so I could scrape Reddit from here"'
)
quote.paragraph_format.left_indent = Inches(0.5)
quote.runs[0].italic = True

doc.add_paragraph(
    'This produced a data collection pipeline (collect_reddit_free.py), an automated schedule runner, '
    'and a five-layer Markov chain predictor (predict_post_flow.py). The initial system could scrape '
    'Reddit data and predict state transitions using baseline transition probabilities computed from '
    'observed state changes.'
)

doc.add_paragraph(
    'Subsequent development used Claude Code throughout, with the developer directing analysis priorities '
    'and evaluating outputs. Additional tools included the VADER sentiment analyser (Hutto and Gilbert, 2014) '
    'for comment scoring and scikit-learn (Pedregosa et al., 2011) for classification and clustering algorithms.'
)

# ===== SECTION 3 =====
doc.add_heading('3. Critique of Initial Code', level=2)

doc.add_paragraph(
    'The initial code contained several technical issues. First, it used datetime.UTC, introduced in '
    'Python 3.11, across twelve files, causing ImportError on the development machine running Python 3.10. '
    'Second, export_history_to_sqlite.py contained backslashes inside f-string expressions, which '
    'Python 3.10 does not permit. Third, the dependency specification required scikit-learn 1.8+, '
    'incompatible with Python 3.10.'
)

doc.add_paragraph(
    'Beyond compatibility, the initial predictor had fundamental algorithmic limitations. It relied '
    'exclusively on upvote and comment velocity, treating a post with 100 angry comments identically '
    'to one with 100 supportive comments. It could only predict one step ahead rather than forecasting '
    'trajectories over hours or days. It used a single global model despite vastly different dynamics '
    'across subreddits; for instance, r/politics posts have a median alive duration of 11 hours compared '
    'to 49 hours for r/Games. The Markov chain approach also produced probability distributions that '
    'converged to a fixed equilibrium regardless of starting conditions, limiting its discriminative '
    'power for long-range prediction.'
)

doc.add_paragraph(
    'The pipeline automation script used WPF MessageBox popups and aggressive error handling that caused '
    'crashes when executed non-interactively by Windows Task Scheduler, rendering the automated collection '
    'system non-functional.'
)

# ===== SECTION 4 =====
doc.add_heading('4. Iterative Development & Justification', level=2)

iterations = [
    ('Iteration 1: Compatibility and Infrastructure Fixes.',
     'Replaced datetime.UTC with timezone.utc across twelve files, extracted problematic f-string '
     'expressions to intermediate variables, and pinned scikit-learn to version 1.7. Fixed the Task '
     'Scheduler script by replacing WPF popups with silent logging and changing the error policy to '
     'Continue. Verified by confirming all scripts executed without errors and the scheduler completed '
     'hourly collections successfully.'),
    ('Iteration 2: VADER Sentiment Integration.',
     'Integrated the VADER sentiment analyser to score all 677,569 comments, adding five features to '
     'the prediction dataset: mean sentiment, upvote-weighted sentiment, positive comment share, negative '
     'comment share, and sentiment variance. This revealed a counterintuitive finding: posts with negative '
     'comment sentiment survive longer (alive posts average sentiment -0.006 versus dead posts +0.045, '
     'difference -0.072). Controversial discussions drive engagement; apathy kills posts. A Decision Tree '
     'classifier using these features achieved 74.5% accuracy (5-fold cross-validation). However, feature '
     'importance analysis revealed comment count dominated at 74%, meaning the model was essentially counting '
     'comments rather than analysing sentiment.'),
    ('Iteration 3: Comment Upvote Gini Coefficient.',
     'To find a stronger engagement signal, implemented Gini coefficient analysis on comment upvote '
     'distributions. The Gini coefficient measures concentration: high values indicate a few dominant '
     'comments capturing most upvotes, while low values indicate evenly distributed attention. This became '
     'the strongest predictor found in the project (46% feature importance). Surviving posts exhibit high '
     'Gini (0.63-0.72), indicating clear community consensus around specific comments. Dying posts show '
     'low Gini (0.36), suggesting diffuse, unfocused discussion. Adding Gini features improved classifier '
     'accuracy to 77.6%.'),
    ('Iteration 4: Per-Subreddit Classification.',
     'Trained separate Random Forest classifiers per subreddit, recognising that each community has '
     'distinct engagement patterns. This improved r/politics accuracy to 81% while revealing that r/Games '
     '(64%) is inherently less predictable. The variation itself is informative: politically focused '
     'communities follow more systematic engagement patterns than entertainment-focused ones.'),
    ('Iteration 5: Multi-Horizon Prediction.',
     'Extended the system from single-step to multi-horizon prediction by training seventeen separate '
     'classifiers for horizons from 1 to 72 hours. Each model was trained on 100,000 to 160,000 labelled '
     'samples. This produced two key outputs: ROC curves showing systematic accuracy decay across horizons '
     '(Figure 1), and survival probability curves estimating post half-lives (Figure 2). The 4-hour model '
     'achieved 0.834 ROC AUC while the 48-hour model degraded to 0.726, quantifying the inherent '
     'unpredictability of long-range Reddit engagement.'),
]

for heading, body in iterations:
    p = doc.add_paragraph()
    p.add_run(heading).bold = True
    p.add_run(' ' + body)

# ===== SECTION 5 =====
doc.add_heading('5. Final Code Evaluation and Reflection', level=2)

p = doc.add_paragraph()
p.add_run('Performance Metrics').bold = True

table = doc.add_table(rows=7, cols=3, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Prediction Target', 'Horizon', 'ROC AUC']):
    cell = table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True

data = [
    ('Post survival', '1 hour', '0.843'),
    ('Post survival', '4 hours', '0.834'),
    ('Post survival', '12 hours', '0.809'),
    ('Post survival', '24 hours', '0.771'),
    ('Post survival', '48 hours', '0.726'),
    ('Surging detection', '1 hour', '0.987'),
]
for i, (t, h, a) in enumerate(data):
    table.rows[i+1].cells[0].text = t
    table.rows[i+1].cells[1].text = h
    table.rows[i+1].cells[2].text = a

doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Survival Probability and Post Half-Life. ').bold = True
p.add_run(
    'By combining predictions across all horizons, the system computes survival probability curves '
    'for any post (Figure 2). Estimated half-lives: surging posts approximately 48 hours, alive posts '
    'approximately 24 hours, cooling posts approximately 3 hours, dying posts approximately 1 hour. '
    'These provide actionable estimates for content monitoring applications.'
)

p = doc.add_paragraph()
p.add_run('Confusion Matrix Analysis. ').bold = True
p.add_run(
    'The 4-hour binary classifier (Figure 3) correctly identifies 80% of surviving posts and 67% of '
    'dying posts. The primary failure mode is false optimism: predicting survival when posts actually '
    'die, reflecting class imbalance in the training data where most snapshots capture posts while '
    'still active.'
)

p = doc.add_paragraph()
p.add_run('Key Findings. ').bold = True
p.add_run(
    'First, controversy drives engagement. Negative comment sentiment correlates with longer post '
    'survival. The most active subreddit (r/politics, average sentiment -0.10) shows strong growth '
    'while the most positive (r/Games, +0.38) is declining. Second, comment upvote concentration '
    'outperforms all other features in predicting survival. Third, predictability decays systematically '
    'over time, suggesting an inherent chaotic component to Reddit engagement beyond approximately '
    '24 hours.'
)

p = doc.add_paragraph()
p.add_run('Limitations. ').bold = True
p.add_run(
    'Reddit\'s public JSON provides limited reply threading data, preventing analysis of argument depth. '
    'The 12-day collection period, while substantial, limits seasonal pattern detection. Regression '
    'models for predicting exact upvote counts performed poorly (R2 = 0.42), confirming that Reddit '
    'virality has a significant random component. All findings are correlational rather than causal.'
)

# ===== SECTION 6 =====
doc.add_heading('6. Reflection on AI-Assisted Coding', level=2)

doc.add_paragraph(
    'Claude Code proved highly effective for rapid prototyping: generating analysis scripts, suggesting '
    'appropriate algorithms, and handling data pipeline boilerplate. The conversational workflow enabled '
    'rapid iteration, moving from verbal description to working code within minutes. The AI correctly '
    'suggested using the Gini coefficient for upvote concentration analysis, which produced the '
    'project\'s strongest finding.'
)

doc.add_paragraph(
    'However, several AI-generated components required significant human correction. The Task Scheduler '
    'script crashed due to unicode em dashes corrupting under Windows cp1252 encoding, a runtime-only '
    'bug invisible in the source code. The AI suggested collecting 10 comments per post hourly without '
    'calculating that this would exceed the 30-minute timeout. The initial sentiment model\'s 74.5% '
    'accuracy appeared promising until feature importance analysis revealed it was essentially counting '
    'comments rather than analysing sentiment, requiring critical human evaluation to identify.'
)

doc.add_paragraph(
    'Validation relied on cross-validation with held-out test sets, feature importance analysis to '
    'ensure models captured meaningful patterns, and confusion matrices to understand specific failure '
    'modes. This multi-layered validation approach was essential because AI-generated code can produce '
    'plausible but misleading results.'
)

doc.add_paragraph(
    'Ethical considerations include responsible scraping practices (1.1-second delays between requests, '
    'compliance with rate limits), analysis of aggregate patterns rather than individual user behaviour, '
    'and transparent documentation of AI tool usage throughout the development process.'
)

# ===== REFERENCES =====
doc.add_heading('References', level=2)

refs = [
    'Bird, S., Klein, E. and Loper, E. (2009) Natural Language Processing with Python. Sebastopol: O\'Reilly Media.',
    'Hutto, C.J. and Gilbert, E.E. (2014) \'VADER: A Parsimonious Rule-based Model for Sentiment Analysis of Social Media Text\', Proceedings of the Eighth International AAAI Conference on Weblogs and Social Media. Ann Arbor, MI, June 2014.',
    'Pedregosa, F. et al. (2011) \'Scikit-learn: Machine Learning in Python\', Journal of Machine Learning Research, 12, pp. 2825-2830.',
]
for ref in refs:
    doc.add_paragraph(ref)

# ===== FIGURES =====
doc.add_page_break()
doc.add_heading('Figures', level=2)

visuals_dir = os.path.join('data', 'analysis', 'reddit', 'visuals')

figures = [
    ('Figure 1', 'roc_prediction_decay.png',
     'ROC curves for post survival prediction across five time horizons (1h, 4h, 12h, 24h, 48h), demonstrating systematic decay in predictive accuracy.'),
    ('Figure 2', 'survival_probability_curves.png',
     'Survival probability curves for five post archetypes with estimated half-life annotations. Surging posts maintain above 50% survival for approximately 48 hours.'),
    ('Figure 3', 'roc_curves_all_horizons.png',
     'Left: ROC curves for all prediction horizons. Right: Confusion matrix for the 4-hour binary classifier showing true positives, true negatives, false positives, and false negatives.'),
    ('Figure 4', 'subreddit_state_mix.png',
     'Distribution of post lifecycle states across all five tracked subreddits.'),
    ('Figure 5', 'flow_trajectory_by_subreddit.png',
     'Post lifecycle flow trajectories by subreddit, illustrating different state transition rates per community.'),
    ('Figure 6', 'live_pulse_dashboard.png',
     'Activity pulse dashboard showing current engagement levels across all monitored subreddits.'),
]

for fig_label, filename, caption in figures:
    p = doc.add_paragraph()
    p.add_run(fig_label).bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    img_path = os.path.join(visuals_dir, filename)
    if os.path.exists(img_path):
        try:
            doc.add_picture(img_path, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            doc.add_paragraph(f'[Image: {filename}]')

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_paragraph()

# Save
output_path = os.path.join('Assessment_2_Report.docx')
doc.save(output_path)
print(f'Saved: {output_path}')
