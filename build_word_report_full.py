"""Build full draft Word document for Assessment 2."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

# ===== TITLE =====
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Predicting Reddit Post Survival Using\nComment Engagement and Sentiment Analysis')
run.bold = True
run.font.size = Pt(16)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Exploring AI: Understanding and Applications (SPC4004)\nAssessment 2 \u2013 Code Generation Project')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('GitHub Repository: ').bold = True
p.add_run('https://github.com/Vasylffff/redit')

# ===== SECTION 1 =====
doc.add_heading('1. Problem Definition & Dataset Justification', level=2)

doc.add_heading('The Problem', level=3)
doc.add_paragraph(
    'Every day, thousands of posts are submitted to Reddit. Some surge to the front page with tens of '
    'thousands of upvotes. Most die within hours with barely any engagement. The question this project '
    'seeks to answer is: can we predict, based on a post\'s early signals, whether it will survive or '
    'die \u2014 and how far into the future can that prediction remain reliable?'
)
doc.add_paragraph(
    'This is framed as a binary classification problem: given a post\'s current state (upvotes, comments, '
    'velocity, age, comment sentiment, and comment engagement patterns), classify it as either surviving '
    '(will be in surging or alive state) or dying (will be in cooling, dying, or dead state) at a future '
    'time horizon.'
)
doc.add_paragraph(
    'The problem extends across multiple time horizons: 1 hour ahead, 4 hours, 12 hours, 24 hours, and '
    '48 hours. This multi-horizon approach allows us to study how predictability itself decays over time '
    '\u2014 a finding that proved to be one of the project\'s most significant contributions.'
)

doc.add_heading('Why This Problem Matters', level=3)
doc.add_paragraph(
    'Understanding post engagement dynamics has practical applications. Content creators want to know which '
    'posts to invest effort in promoting. News monitoring systems need to identify which stories are gaining '
    'traction. Researchers studying online discourse benefit from understanding what drives engagement versus '
    'apathy. From a machine learning perspective, it presents an interesting challenge: a time-series '
    'classification problem with naturally labelled data, high dimensionality, and an inherent stochastic component.'
)

doc.add_heading('The Dataset', level=3)
doc.add_paragraph(
    'The dataset was entirely self-collected using Reddit\'s public JSON endpoints (for example, '
    'https://www.reddit.com/r/technology/new.json). No API key or authentication was required. A custom '
    'scraping system was built that runs on Windows Task Scheduler, collecting data hourly across five '
    'subreddits: r/technology, r/news, r/worldnews, r/politics, and r/Games.'
)
doc.add_paragraph('Collection period: 12 days (26 March to 6 April 2026).')

table = doc.add_table(rows=5, cols=2, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (m, v) in enumerate([('Metric', 'Value'), ('Post snapshots', '185,048'), ('Comment snapshots', '677,569'), ('Unique post lifecycles', '6,292'), ('Subreddit health observations', '2,174')]):
    table.rows[i].cells[0].text = m
    table.rows[i].cells[1].text = v
    if i == 0:
        for cell in table.rows[i].cells:
            cell.paragraphs[0].runs[0].bold = True

doc.add_paragraph()
doc.add_paragraph(
    'Each post snapshot records: upvote count, upvote ratio, comment count, upvote velocity, comment velocity, '
    'age in minutes, rank position, and computed activity state. Comment snapshots include: full text body, '
    'comment upvote count, author, and whether it is top-level or a reply.'
)
doc.add_paragraph(
    'This dataset was chosen because it provides naturally labelled training data: lifecycle states serve as '
    'ground truth without manual annotation. The time-series structure enables trajectory prediction, and '
    'comment text enables sentiment and engagement analysis.'
)

doc.add_heading('Activity States', level=3)
for state, desc in [('Surging', 'Velocity significantly above subreddit median.'), ('Alive', 'Velocity above minimum threshold.'), ('Cooling', 'Velocity declining but above zero.'), ('Dying', 'Velocity near zero.'), ('Dead', 'No meaningful engagement across multiple snapshots.')]:
    p = doc.add_paragraph()
    p.add_run(state + ': ').bold = True
    p.add_run(desc)

# ===== SECTION 2 =====
doc.add_heading('2. Initial Code & Explanation of AI Use', level=2)
doc.add_paragraph('The entire initial codebase was generated using Claude Code, Anthropic\'s CLI powered by Claude Opus 4.6. The first prompt was:')
quote = doc.add_paragraph('"This is my project and I want you to create a folder and set this up so I could scrape Reddit from here"')
quote.paragraph_format.left_indent = Inches(0.5)
quote.runs[0].italic = True

doc.add_paragraph('This produced:')
for item in ['collect_reddit_free.py \u2014 Reddit scraper with 1.1s delay', 'run_free_collection_schedule.py \u2014 schedule runner', 'run_free_collection_window.ps1 \u2014 PowerShell Task Scheduler wrapper', 'build_reddit_history.py \u2014 raw JSON to structured CSV', 'predict_post_flow.py \u2014 five-layer Markov chain predictor']:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('Additional tools: VADER sentiment analyser (Hutto and Gilbert, 2014) and scikit-learn (Pedregosa et al., 2011).')

# ===== SECTION 3 =====
doc.add_heading('3. Critique of Initial Code', level=2)
doc.add_heading('Correctness Issues', level=3)
for h, d in [('Python 3.11 incompatibility.', 'Used datetime.UTC across 12 files; fails on Python 3.10.'), ('f-string syntax error.', 'Backslashes in f-strings, not permitted in Python 3.10.'), ('Dependency mismatch.', 'scikit-learn>=1.8.0 requires Python 3.11+.')]:
    p = doc.add_paragraph()
    p.add_run(h).bold = True
    p.add_run(' ' + d)

doc.add_heading('Algorithmic Limitations', level=3)
for h, d in [('No sentiment analysis.', 'Blind to emotional content of discussions.'), ('No comment engagement analysis.', 'Ignored upvote distribution patterns (Gini coefficient later proved strongest predictor).'), ('Single-step prediction.', 'Markov chain converges to equilibrium; cannot discriminate at longer horizons.'), ('No per-subreddit models.', 'r/politics (11h alive) treated same as r/Games (49h alive).')]:
    p = doc.add_paragraph()
    p.add_run(h).bold = True
    p.add_run(' ' + d)

doc.add_heading('Infrastructure Issues', level=3)
doc.add_paragraph('Task Scheduler crashed due to WPF popup dependency in non-interactive sessions and aggressive error handling treating Python stderr as fatal errors.')

# ===== SECTION 4 =====
doc.add_heading('4. Iterative Development & Justification', level=2)

iters = [
    ('Iteration 1: Compatibility and Infrastructure Fixes',
     [('Problem:', ' Entire codebase non-functional on Python 3.10; Task Scheduler crashed silently.'),
      ('Changes:', ' Fixed datetime.UTC in 12 files, f-string syntax, scikit-learn version. Removed WPF popups, fixed error handling, replaced unicode em dashes.'),
      ('Verification:', ' All scripts run clean. Scheduler completes hourly with exit code 0.')]),
    ('Iteration 2: VADER Sentiment Integration',
     [('Problem:', ' Model blind to comment sentiment.'),
      ('Changes:', ' Scored 677,569 comments with VADER. Added 5 sentiment columns. Added Layer 6 to Markov predictor.'),
      ('Key finding:', ' Negative sentiment correlates with longer survival (alive avg -0.006 vs dead +0.045). Controversy drives engagement. Decision Tree achieved 74.5% accuracy, but comment count dominated at 74% feature importance \u2014 model was counting comments, not analysing sentiment.'),
      ('Verification:', ' Consistent across all 5 subreddits and 3,520 posts.')]),
    ('Iteration 3: Comment Upvote Gini Coefficient',
     [('Problem:', ' Sentiment added only 2-3% above comment volume.'),
      ('Changes:', ' Implemented Gini coefficient on comment upvote distributions. High Gini = few dominant comments (consensus); low Gini = diffuse attention (unfocused).'),
      ('Key finding:', ' Gini became #1 predictor (46% importance). Surging/alive: 0.63-0.72. Dying: 0.36. Dying posts also showed largest sentiment gap (-0.135): community endorsed angry comments. Classifier improved to 77.6%.'),
      ('Why this matters:', ' Survival depends not on whether people are happy or angry, but whether discussion has focus.'),
      ('Verification:', ' 5-fold CV. Pattern held across all subreddits.')]),
    ('Iteration 4: Per-Subreddit Classification',
     [('Problem:', ' Global model ignores community-specific dynamics.'),
      ('Changes:', ' Separate Random Forest per subreddit.'),
      ('Results:', ' r/politics 81%, r/worldnews 76%, r/technology 70%, r/news 68%, r/Games 64%.'),
      ('Verification:', ' Distinct confusion matrix patterns per community.')]),
    ('Iteration 5: Multi-Horizon Prediction',
     [('Problem:', ' Only single-step prediction.'),
      ('Changes:', ' 17 classifiers for 1-72h horizons, each trained on 100K-160K samples. Generated ROC curves, survival probability curves, half-life estimates.'),
      ('Results:', ' ROC AUC: 1h=0.843, 4h=0.834, 12h=0.809, 24h=0.771, 48h=0.726. Half-lives: surging ~48h, alive ~24h, cooling ~3h, dying ~1h.'),
      ('Significance:', ' Demonstrates deterministic component (first 4-12h) and stochastic component (grows over time). Boundary at ~24 hours.'),
      ('Verification:', ' 5-fold CV at each horizon. Smooth monotonic decay curve.')]),
]

for title_text, paragraphs in iters:
    doc.add_heading(title_text, level=3)
    for label, text in paragraphs:
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        p.add_run(text)

doc.add_heading('Additional Analyses', level=3)
for item in ['Best posting hours (08:00 UTC optimal)', 'Upvote velocity curves by subreddit', 'Cross-subreddit propagation (1,573 stories, 11.1h median)', 'Subreddit direction forecasting', 'Domain and author performance analysis', 'Keyword trend detection', 'Title style impact (shock words: 59.1% alive vs 23.8%)', 'Cross-posting success predictor (politics\u2192news: 72%)', 'Time-to-death regression (R2=0.459, MAE 9.1h)']:
    doc.add_paragraph(item, style='List Bullet')

# ===== SECTION 5 =====
doc.add_heading('5. Final Code Evaluation and Reflection', level=2)

p = doc.add_paragraph()
p.add_run('Classification Performance:').bold = True
t = doc.add_table(rows=7, cols=3, style='Light Shading Accent 1')
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Model', 'Horizon', 'ROC AUC']):
    t.rows[0].cells[i].text = h
    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
for i, (m, h, a) in enumerate([('Post survival','1 hour','0.843'),('Post survival','4 hours','0.834'),('Post survival','12 hours','0.809'),('Post survival','24 hours','0.771'),('Post survival','48 hours','0.726'),('Surging detection','1 hour','0.987')]):
    t.rows[i+1].cells[0].text = m; t.rows[i+1].cells[1].text = h; t.rows[i+1].cells[2].text = a
doc.add_paragraph()

doc.add_heading('Confusion Matrix (4-Hour)', level=3)
cm = doc.add_table(rows=3, cols=3, style='Light Shading Accent 1')
cm.alignment = WD_TABLE_ALIGNMENT.CENTER
cm.rows[0].cells[0].text = ''; cm.rows[0].cells[1].text = 'Predicted Alive'; cm.rows[0].cells[2].text = 'Predicted Dead'
cm.rows[1].cells[0].text = 'Actually Alive'; cm.rows[1].cells[1].text = '16,570 (TP)'; cm.rows[1].cells[2].text = '4,189 (FN)'
cm.rows[2].cells[0].text = 'Actually Dead'; cm.rows[2].cells[1].text = '3,900 (FP)'; cm.rows[2].cells[2].text = '8,057 (TN)'
for cell in cm.rows[0].cells: cell.paragraphs[0].runs[0].bold = True
for row in cm.rows[1:]: row.cells[0].paragraphs[0].runs[0].bold = True
doc.add_paragraph('80% true positive rate, 67% true negative rate. Primary failure: false optimism from class imbalance.')

doc.add_heading('Key Findings', level=3)
for h, d in [('1. Controversy drives engagement.', 'Negative sentiment correlates with survival. r/politics (-0.10 sentiment) grows; r/Games (+0.38) declines.'), ('2. Upvote concentration is the strongest predictor.', 'Gini coefficient outperforms all other features.'), ('3. Predictability decays over time.', 'ROC 0.843 (1h) to 0.726 (48h). Boundary at ~24 hours.'), ('4. Half-life varies by state.', 'Surging ~48h, alive ~24h, cooling ~3h, dying ~1h.'), ('5. Per-subreddit models outperform global.', '+10 percentage points for r/politics.')]:
    p = doc.add_paragraph()
    p.add_run(h).bold = True
    p.add_run(' ' + d)

doc.add_heading('Limitations', level=3)
doc.add_paragraph('Reply threading data mostly empty (99% zeros). 12-day window limits seasonal detection. Regression models weak (R2=0.42). All correlational, not causal.')

# ===== SECTION 6 =====
doc.add_heading('6. Reflection on AI-Assisted Coding', level=2)
doc.add_heading('Effective Uses', level=3)
doc.add_paragraph('Rapid prototyping, algorithm selection (Random Forest, VADER, K-means), boilerplate generation, and novel suggestions (Gini coefficient, multi-horizon prediction).')

doc.add_heading('Corrections Required', level=3)
for h, d in [('Task Scheduler crash:', ' WPF + unicode em dashes; works interactively, crashes in deployment.'), ('Timeout miscalculation:', ' 5\u00d7100\u00d710 comments\u00d71.1s = 92min, exceeding 30min limit.'), ('Misleading accuracy:', ' 74.5% appeared good until feature importance showed 74% was just comment counting.'), ('Version assumptions:', ' Generated Python 3.11 code for 3.10 machine.')]:
    p = doc.add_paragraph()
    p.add_run(h).bold = True
    p.add_run(d)

doc.add_heading('Validation', level=3)
doc.add_paragraph('Execution testing, 5-fold cross-validation, feature importance analysis, confusion matrices, and cross-subreddit sanity checks on counterintuitive findings.')

doc.add_heading('Ethics', level=3)
doc.add_paragraph('1.1-second rate limiting, aggregate analysis not individual profiling, transparent AI documentation.')

# ===== REFERENCES =====
doc.add_heading('References', level=2)
for ref in ['Bird, S., Klein, E. and Loper, E. (2009) Natural Language Processing with Python. O\'Reilly Media.', 'Hutto, C.J. and Gilbert, E.E. (2014) \'VADER: A Parsimonious Rule-based Model for Sentiment Analysis of Social Media Text\', Proc. ICWSM.', 'Pedregosa, F. et al. (2011) \'Scikit-learn: Machine Learning in Python\', JMLR, 12, pp. 2825-2830.']:
    doc.add_paragraph(ref)

# ===== FIGURES =====
doc.add_page_break()
doc.add_heading('Figures', level=2)
visuals = 'data/analysis/reddit/visuals'
for label, fname, cap in [
    ('Figure 1: ROC Prediction Decay', 'roc_prediction_decay.png', 'ROC curves across 5 horizons showing predictability decay from AUC=0.843 (1h) to 0.726 (48h).'),
    ('Figure 2: Survival Probability Curves', 'survival_probability_curves.png', 'Survival probability for 5 post types over 72 hours with half-life annotations.'),
    ('Figure 3: ROC Curves + Confusion Matrix', 'roc_curves_all_horizons.png', 'All ROC curves (left) and 4-hour confusion matrix (right).'),
    ('Figure 4: Subreddit State Distribution', 'subreddit_state_mix.png', 'Lifecycle state distribution across subreddits.'),
    ('Figure 5: Flow Trajectories', 'flow_trajectory_by_subreddit.png', 'State transition patterns per subreddit.'),
    ('Figure 6: Activity Dashboard', 'live_pulse_dashboard.png', 'Current engagement levels across subreddits.'),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(label).bold = True
    path = os.path.join(visuals, fname)
    if os.path.exists(path):
        try:
            doc.add_picture(path, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except: doc.add_paragraph(f'[{fname}]')
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(cap)
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(100,100,100)
    doc.add_paragraph()

doc.save('Assessment_2_Full_Draft.docx')
print('Saved: Assessment_2_Full_Draft.docx')
