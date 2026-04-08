"""Build final Assessment 2 Word document with complete timeline and lifecycle focus."""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

FIGURES_DIR = "data/analysis/reddit/figures"
OUTPUT_PATH = "Assessment_2_Final_v2.docx"

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles["Heading %d" % level]
    hs.font.name = "Calibri"
    hs.font.color.rgb = RGBColor(0, 0, 0)


def para(text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(text)
    run.font.bold = bold
    run.font.italic = italic
    return p


def quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(text)
    run.font.italic = True
    return p


def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")


def add_table(headers, rows_data):
    table = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows_data):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()


def fig(name, caption):
    fpath = os.path.join(FIGURES_DIR, name)
    if os.path.exists(fpath):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(fpath, width=Inches(5.5))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        run = cap.add_run(caption)
        run.font.size = Pt(9)
        run.font.italic = True


# ================================================================
# TITLE PAGE
# ================================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Predicting Reddit Topic Lifecycles:\nEmergence, Spread, and Death")
run.font.size = Pt(18)
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Exploring AI: Understanding and Applications (SPC4004)\nAssessment 2 -- Code Generation Project\n\nVasyl Shcherbatykh")
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("GitHub: https://github.com/Vasylffff/redit")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0, 102, 204)

doc.add_page_break()

# ================================================================
# 1. PROBLEM DEFINITION
# ================================================================
doc.add_heading("1. Problem Definition & Dataset Justification", level=1)

doc.add_heading("What This Project Does", level=2)

para(
    "Every day, thousands of stories appear on Reddit. Some disappear within hours. Others spread "
    "across multiple communities, generate thousands of comments, and dominate the platform for days. "
    "This project asks: can we predict which stories will grow, how they will spread, and when they "
    "will die, using only early engagement signals?"
)

para(
    "The system tracks topics through their complete lifecycle: emergence (a new story appears in "
    "1-3 posts), growth (it gains momentum), spread (it crosses into new subreddits), decline "
    "(engagement drops), and death (the story stops generating new posts). At each stage, a separate "
    "machine learning model makes predictions about what will happen next."
)

doc.add_heading("The Machine Learning Tasks", level=2)

para(
    "The project applies multiple types of machine learning to different lifecycle stages:"
)

bullet("Binary classification: given a word pair with 1-3 posts today, will it reach 5+ posts "
       "within 3 days? Input features: post count, total upvotes, subreddit count, best comment "
       "count, average comments, upvotes per post, best velocity, average velocity. Output: grow (1) "
       "or not (0).")
bullet("Multi-class classification: given a topic's current state (surging, growing, stable, "
       "cooling, dying, dead), what state will it be in tomorrow? Input: current engagement metrics "
       "plus previous state. Output: one of six states.")
bullet("Binary classification with temporal features: given a topic that has just dropped below "
       "its activity threshold, will it revive (ongoing story) or stay dead (one-shot event)? "
       "Input: peak size, number of previous peaks, consistency, active days. Output: revive (1) "
       "or dead (0).")
bullet("Regression (attempted): given a topic at 1-3 posts, how many posts will it reach at peak? "
       "This task was attempted using multiple approaches including Szabo-Huberman log-linear models "
       "(2010) but did not achieve useful accuracy (R-squared = 0.22), which itself became a finding.")

para(
    "Five classifiers were compared across seven tasks: Random Forest, Extra Trees, Gradient "
    "Boosting, Logistic Regression, and Decision Tree, with 36 hyperparameter configurations tested."
)

doc.add_heading("Why This Problem Matters", level=2)

para(
    "Understanding how topics emerge and die online has practical applications in several areas. "
    "News monitoring systems need to identify which stories are gaining traction before they peak. "
    "Content moderators need to anticipate which topics will spread across communities. Researchers "
    "studying misinformation need to understand how stories propagate between subreddits and whether "
    "spread patterns can be detected early."
)

para(
    "From a machine learning perspective, the problem is interesting because it requires prediction "
    "at different timescales and granularities. Post-level prediction (will this individual post "
    "survive?) works well at short horizons but degrades beyond 24 hours. Topic-level prediction "
    "(will this story grow?) follows the opposite pattern: it improves at longer horizons because "
    "topics follow momentum while individual posts are chaotic. This inverse relationship between "
    "post and topic predictability was one of the project's key discoveries."
)

doc.add_heading("Why This Dataset", level=2)

para(
    "The system operates at two levels. At the post level, it incorporates text-based features "
    "including VADER sentiment scoring of 972,353 comments and comment engagement patterns such as "
    "the Gini coefficient of upvote distribution. At the topic level, the system is content-agnostic: "
    "it tracks co-occurring word pairs in titles purely through numerical engagement signals such as "
    "post counts, upvotes, and velocity, without interpreting what those words mean. This design was "
    "motivated by the observation that viral content of any type, whether political news, game "
    "announcements, or internet memes, follows similar spread patterns driven by human attention "
    "rather than content meaning."
)

para(
    "Five subreddits were chosen to represent different community types: r/politics and r/news "
    "(fast-moving political discussion), r/worldnews (international affairs), r/technology (industry "
    "news), and r/Games (entertainment). This diversity tests whether the same models generalise "
    "across communities with different engagement patterns. For example, r/politics posts have an "
    "average lifecycle of 11 hours while r/Games posts survive 49 hours."
)

para(
    "The dataset was self-collected rather than using an existing corpus because the project requires "
    "repeated observations of the same posts over time. Existing Reddit datasets typically provide "
    "single snapshots. Hourly collection over 13 days produced 216,944 post snapshots capturing how "
    "each post's engagement changed hour by hour, which is essential for velocity and trajectory features."
)

doc.add_heading("Dataset", level=2)

para(
    "The dataset was entirely self-collected over 13 days (26 March to 7 April 2026) using Reddit's "
    "public JSON endpoints. Five subreddits were tracked: r/technology, r/news, r/worldnews, r/politics, "
    "and r/Games. Data was collected hourly via Windows Task Scheduler running on two machines, with "
    "results merged to improve coverage."
)

add_table(
    ["Metric", "Value"],
    [
        ["Post snapshots", "216,944"],
        ["Comment snapshots", "972,353"],
        ["Unique posts tracked", "6,826"],
        ["Collection frequency", "Hourly"],
        ["Raw JSON files", "5,140"],
        ["Co-occurrence pairs analysed", "200,000+ per day"],
    ]
)

para(
    "Posts were assigned lifecycle states (surging, alive, cooling, dying, dead) based on upvote "
    "and comment velocity thresholds computed per subreddit, providing natural labels without manual "
    "annotation. Topics were represented as two-word pairs extracted from titles, following the "
    "co-occurrence approach described by Szabo and Huberman (2010) for tracking content popularity."
)

fig("fig11_data_coverage.png", "Figure 1: Daily data collection coverage by subreddit across the 13-day collection period.")

# ================================================================
# 2. INITIAL CODE & AI USE
# ================================================================
doc.add_heading("2. Initial Code & Explanation of AI Use", level=1)

para(
    "The project was developed using three AI tools across multiple sessions. Development began on "
    "24 March 2026 with OpenAI Codex, where the initial prompt was:"
)

quote('"ok i want to make a project that will analyse redit flow and predict it. Firstly i wnat to take a new data"')

para(
    "This session produced the initial data collection pipeline, including Reddit API access attempts "
    "via PRAW, an Apify-based fallback, and ultimately a free public JSON endpoint scraper. The session "
    "spanned 352 prompts over 7 days (24-31 March) and established the core architecture: hourly "
    "collection, history building, gap patching, and prediction dataset generation."
)

para(
    "A second Codex session (31 March) performed a major pipeline refactoring documented in a formal "
    "handover summary. This session separated prediction tracking from live monitoring, redesigned "
    "the lifecycle state model, and fixed critical scheduler bugs."
)

para(
    "Claude Code (Anthropic, Claude Opus 4.6) was used from 31 March onwards for analysis and "
    "modelling. This included sentiment analysis, comment engagement features, multi-horizon prediction, "
    "and the complete topic lifecycle pipeline developed on 7-8 April."
)

para(
    "All subsequent development was conversational: the developer directed priorities, questioned "
    "results, and pushed exploration when AI suggested premature ceilings. The choice to pursue "
    "topic lifecycle prediction rather than stopping at post-level models was entirely developer-driven."
)

# ================================================================
# 3. CRITIQUE OF INITIAL CODE
# ================================================================
doc.add_heading("3. Critique of Initial Code", level=1)

doc.add_heading("Bugs", level=2)

bullet("datetime.UTC attribute incompatible with Python 3.10, present across 12 files")
bullet("f-string backslash syntax errors causing runtime failures")
bullet("scikit-learn version incompatibility requiring pinning to 1.7+")
bullet("Windows Task Scheduler crashed in non-interactive mode due to WPF popup windows")
bullet("Unicode em dash characters corrupted under Windows cp1252 encoding")
bullet("CSV headers contained hidden UTF-8 BOM characters, breaking manifest validation")

doc.add_heading("Algorithmic Limitations", level=2)

para(
    "No topic-level analysis. The initial system predicted individual post trajectories using a "
    "five-layer Markov chain, but had no mechanism to detect that multiple posts referred to the "
    "same underlying story. The concept of a \"topic\" was entirely absent from the original codebase."
)

para(
    "Single-keyword vocabulary. When topic analysis was first introduced, it used individual "
    "keywords (\"trump\", \"iran\"). This proved too generic: the keyword \"trump\" appeared in 2,192 "
    "posts and carried no discriminative signal. The transition to co-occurrence pairs was a critical "
    "improvement that the AI did not initially suggest."
)

para(
    "Uniform model selection. Random Forest was applied to every task with near-default hyperparameters. "
    "Subsequent testing across five classifiers and 36 parameter configurations revealed that Logistic "
    "Regression outperforms Random Forest on the core emergence detection task (0.860 vs 0.829 ROC AUC), "
    "indicating the underlying signal is linear rather than requiring complex feature interactions."
)

para(
    "Incorrect topic death definition. The initial code classified a topic as dead after a single day "
    "below two posts. Analysis of 8,880 drop events revealed a 13.1% false-death rate, with 44.8% of "
    "topics peaking at 8-11 posts reviving after apparent death. A two-consecutive-day threshold "
    "reduced false deaths to 6.8%."
)

fig("fig4_depth_vs_roc.png",
    "Figure 2: Overfitting analysis. Both Random Forest and Decision Tree perform worse with deeper trees. "
    "Optimal depth is 3-5, not the default 8-12.")

# ================================================================
# 4. ITERATIVE DEVELOPMENT
# ================================================================
doc.add_heading("4. Iterative Development & Justification", level=1)

doc.add_heading("Phase 1: Data Collection (March 24-26)", level=2)

para(
    "Iteration 1: API access and collection pivots. The initial approach required Reddit API "
    "credentials via PRAW. The API application was submitted on 24 March but not approved. Apify "
    "(a commercial web scraping service) was adopted as a fallback the same day. On 26 March, "
    "the developer discovered that Reddit's public JSON endpoints (appending .json to any listing URL) "
    "provide post metadata without authentication. This became the primary collection method, "
    "reducing cost to zero and removing the API dependency. Collection was automated via Windows "
    "Task Scheduler with 1.1-second rate limiting between requests."
)

para(
    "Iteration 2: Automated hourly collection. A schedule system was built with five cadences: "
    "hourly (new posts), every 2 hours (rising), every 4 hours (hot), twice daily (top/day), "
    "and daily (top/week). Task Scheduler was installed on two machines to improve coverage, and "
    "results were merged to fill gaps caused by laptop sleep and subway commutes."
)

doc.add_heading("Phase 2: Post-Level Prediction (March 27 - April 1)", level=2)

para(
    "Iteration 3: History building and gap patching. The build_reddit_history.py script merged "
    "all raw JSON snapshots into a unified timeline. patch_snapshot_gaps.py addressed velocity "
    "corruption caused by missed collection windows, where a gap would incorrectly show zero "
    "velocity even for actively growing posts. Gaps under 3 hours were interpolated; longer gaps "
    "were flagged but velocity was still recalculated from the upvote delta."
)

para(
    "Iteration 4: VADER sentiment analysis. 972,353 comments were scored using VADER. The finding "
    "was counterintuitive: negative comment sentiment correlates with longer post survival (alive "
    "average -0.006 vs dead +0.045). Controversy drives engagement; apathy kills posts."
)

para(
    "Iteration 5: Comment engagement features. The Gini coefficient of comment upvote distribution "
    "was introduced to measure whether discussion is concentrated (a few dominant comments) or "
    "diffuse. This became the strongest single post-level predictor at 46% feature importance."
)

para(
    "Iteration 6: Multi-horizon classifiers. Seventeen classifiers were built for prediction "
    "horizons from 1 hour to 7 days. ROC AUC decayed from 0.843 (1 hour) to approximately 0.57 "
    "(7 days). At 7 days, accuracy paradoxically rose to 85% while ROC dropped, because predicting "
    "\"everything dies\" achieves high accuracy but zero discriminative power. This demonstrated "
    "why ROC AUC is a more appropriate metric than accuracy for imbalanced classification."
)

para(
    "Iteration 7: Pipeline refactoring (Codex session 2). The tracking system was split into "
    "two lanes: a fixed cohort for prediction research and a rolling shortlist for live monitoring. "
    "The lifecycle state model was redesigned to provide earlier warning signals through the "
    "\"dying\" state. Scheduler bugs involving BOM characters in CSV headers were resolved."
)

doc.add_heading("Phase 3: Topic Lifecycle Prediction (April 7-8)", level=2)

para(
    "Iteration 8: Co-occurrence pair detection. The transition from single keywords to two-word "
    "pairs was the most significant methodological improvement. Pairs such as \"russian+tanker\" "
    "and \"birthright+citizenship\" represent specific stories rather than generic vocabulary. "
    "Temporal validation (trained on days 1-8, tested on days 9-13) achieved 0.813 ROC AUC for "
    "detecting topics at 1-3 posts that would grow to 5 or more."
)

fig("fig1_topic_trajectories.png",
    "Figure 3: Topic lifecycle trajectories across 13 days. Ongoing stories (Hormuz Strait) show "
    "repeated surges and recoveries. One-shot events (Easter+Trump) spike once and die.")

para(
    "Iteration 9: Content-agnostic detection. The same algorithm detected entirely different "
    "content types: \"official+trailer\" (122 posts, game announcements), \"hormuz+strait\" (252 posts, "
    "geopolitical crisis), \"crimson+desert\" (65 posts, game launch), and \"media+social\" (127 posts, "
    "technology policy). The model operates on engagement patterns alone and does not process "
    "text content. A political scandal and a viral game trailer follow identical spread dynamics "
    "from the classifier's perspective."
)

para(
    "Iteration 10: Growth peak detection. A classifier was trained to predict whether a topic "
    "had already reached its peak or would continue growing, using the first two days of data. "
    "ROC AUC: 0.958. The dominant features were post count growth between day 1 and day 2 "
    "(52.7% importance) and upvote growth (38.1%)."
)

para(
    "Iteration 11: Topic death prediction. A model to predict whether a topic would die the "
    "following day achieved 0.890 ROC AUC. Current post count (46.1%) and subreddit coverage "
    "(17.1%) were the dominant features. Losing subreddit presence is a strong signal of imminent death."
)

para(
    "Iteration 12: Death speed classification. After a topic peaks, will it die within 0-1 days "
    "(quick death) or survive 2+ days (slow death)? ROC AUC: 0.996. The decisive feature was "
    "the decline rate on the first day after the peak (67.8% importance). A sharp drop signals "
    "rapid death; gradual decline indicates an ongoing story."
)

fig("fig5_pipeline_summary.png",
    "Figure 4: Complete topic lifecycle prediction pipeline. Green bars indicate ROC AUC above 0.9; "
    "orange indicates above 0.8.")

para(
    "Iteration 13: Subreddit spread prediction. A per-target classifier predicted whether a topic "
    "would appear in a new subreddit the following day. Predicting spread to r/politics achieved "
    "0.756 ROC AUC. Analysis of spread patterns revealed that r/politics breaks stories first "
    "(47,580 times), contrary to the expectation that r/news would lead. The most common spread "
    "route was news to worldnews (2,354 instances)."
)

fig("fig8_subreddit_spread.png",
    "Figure 5: Left: which subreddit breaks stories first. Right: most common cross-subreddit spread routes.")

para(
    "Iteration 14: Ongoing story vs one-shot event classification. When a topic drops below "
    "its activity threshold, the system predicts whether it will revive (ongoing story) or remain "
    "dead (one-shot event). ROC AUC: 0.970. The strongest feature was the number of previous "
    "peaks (21.5% importance). Ongoing stories exhibit multiple surges and dips; one-shot events "
    "peak once."
)

add_table(
    ["Signal", "Ongoing story", "One-shot event", "Ratio"],
    [
        ["Multiple peaks", "0.9", "0.1", "15.5x"],
        ["Consistency (active/total days)", "0.4", "0.1", "4.6x"],
        ["Active days", "3.2", "0.8", "4.0x"],
        ["Peak posts", "3.8", "1.6", "2.3x"],
    ]
)

fig("fig12_ongoing_vs_oneshot.png",
    "Figure 6: Feature ratios distinguishing ongoing stories from one-shot events. "
    "The number of previous peaks is the strongest signal at 15.5x.")

para(
    "Iteration 15: Topic death definition analysis. The original one-day death definition "
    "was found to have a 13.1% false-death rate. Topics with higher peak post counts revived "
    "more frequently: 44.8% for topics peaking at 8-11 posts versus 22.8% for topics peaking "
    "at 3-4 posts. Adopting a two-consecutive-day definition reduced false deaths to 6.8%."
)

fig("fig6_revival_rates.png",
    "Figure 7: Left: false death rates by definition strictness. Right: revival rates by topic size. "
    "Larger topics are substantially more likely to revive.")

para(
    "Iteration 16: Noise filtering. A classifier trained to identify topics that would fail to "
    "grow achieved 0.824 ROC AUC. At 99% confidence, 87% of all word pairs could be discarded "
    "as noise with 99.5% precision, reducing the monitoring set from approximately 36,000 to 5,000 "
    "candidates per day."
)

para(
    "Iteration 17: Model comparison and hyperparameter tuning. Five classifiers were compared "
    "across all seven lifecycle tasks. No single model dominated."
)

fig("fig7_model_heatmap.png",
    "Figure 8: Model performance across all tasks. Bold values indicate the best model per task. "
    "Logistic Regression wins on emergence detection; Random Forest wins on death prediction.")

add_table(
    ["Model", "Default ROC", "Tuned ROC", "Improvement"],
    [
        ["Decision Tree", "0.577", "0.842", "+0.265"],
        ["Gradient Boosting", "0.713", "0.835", "+0.122"],
        ["Random Forest", "0.820", "0.846", "+0.027"],
        ["Logistic Regression", "0.859", "0.860", "+0.001"],
    ]
)

fig("fig10_gbm_learning_rate.png",
    "Figure 9: Gradient Boosting learning rate tuning. Performance peaks at 0.02 and degrades "
    "substantially at higher rates due to overfitting.")

# ================================================================
# 5. FINAL EVALUATION
# ================================================================
doc.add_heading("5. Final Code Evaluation and Reflection", level=1)

doc.add_heading("Post-Level Prediction Results", level=2)

para(
    "Post-level models predict the trajectory of individual posts. These were developed "
    "first and provided the foundation for topic-level work."
)

add_table(
    ["Task", "ROC AUC", "Key Feature"],
    [
        ["Surging detection", "0.987", "Upvote velocity"],
        ["State rise prediction (1h)", "0.947", "Current velocity + comment rate"],
        ["Dead detection", "0.945", "Velocity collapse"],
        ["Survival (1h)", "0.843", "Gini coefficient (46%)"],
        ["Survival (4h)", "0.834", "Gini + velocity"],
        ["Survival (24h)", "0.771", "Per-subreddit models"],
        ["Survival (7d)", "~0.57", "Base rate dominates"],
    ]
)

para(
    "Post prediction accuracy decays with time horizon: from 0.843 ROC at 1 hour to approximately "
    "0.57 at 7 days. At longer horizons, most posts are dead, so predicting \"everything dies\" "
    "achieves 85% accuracy but 0.57 ROC, demonstrating zero discriminative power. Per-subreddit "
    "models improved results: r/politics achieved 81% accuracy while r/Games reached only 64%, "
    "reflecting the more systematic engagement patterns of political communities."
)

para(
    "The Gini coefficient of comment upvote distribution was the strongest single feature at "
    "46% importance. High Gini (0.63-0.72) indicates a few comments dominating the discussion "
    "(community consensus), which predicts survival. Low Gini (0.36) indicates diffuse attention "
    "and predicts death. VADER sentiment analysis of 972,353 comments revealed that negative "
    "sentiment correlates with longer survival (alive average -0.006 vs dead +0.045), confirming "
    "that controversy drives engagement."
)

doc.add_heading("Topic Lifecycle Pipeline Results", level=2)

para(
    "Topic-level models predict the trajectory of stories (represented as co-occurrence word pairs) "
    "across their complete lifecycle. These were developed after post-level models reached their limits."
)

add_table(
    ["Lifecycle Stage", "Task", "Best ROC AUC", "Best Model"],
    [
        ["Filter", "Discard noise (87% filtered)", "0.850", "Logistic Regression"],
        ["Birth", "Detect emerging topic (1-3 to 5+ posts)", "0.860", "Logistic Regression"],
        ["Growth", "Has it peaked or still growing?", "0.958", "Gradient Boosting"],
        ["Spread", "Will it reach r/politics?", "0.756", "Random Forest"],
        ["Decline", "Topic dying state detection", "0.992", "Random Forest"],
        ["Death", "Will it die tomorrow?", "0.890", "Random Forest"],
        ["Death speed", "Quick death or slow death?", "0.999", "Logistic Regression"],
        ["Revival", "Ongoing story or one-shot?", "0.970", "Random Forest"],
    ]
)

para(
    "A key discovery was the inverse relationship between post and topic predictability. "
    "Individual posts become chaotic beyond 24 hours (ROC decays from 0.843 to 0.57). "
    "Topics become more predictable at longer horizons because they follow momentum: "
    "a story with 5 posts today is likely to have more tomorrow, regardless of what happens "
    "to any individual post. r/politics has the best post prediction (81%) but the worst topic "
    "prediction (R-squared = 0.551). r/news is the reverse: worst posts (68%) but best topics "
    "(R-squared = 0.808)."
)

fig("fig2_roc_emergence_models.png",
    "Figure 10: ROC curves comparing five classifiers on topic emergence detection.")

doc.add_heading("Topic State Transitions", level=2)

fig("fig13_transition_matrix.png",
    "Figure 11: Topic state transition matrix. Surging topics have only a 4% chance of remaining "
    "surging the next day. Dead topics stay dead 94% of the time, but 6% revive.")

para(
    "The transition matrix reveals that topic states are highly transient. Surging topics most "
    "commonly transition to stable (48%) or dead (30%). Notably, dying topics revive to stable "
    "44% of the time, confirming that topic death is not permanent and reinforcing the need for "
    "the ongoing-vs-one-shot classifier."
)

doc.add_heading("What Does Not Work", level=2)

para(
    "Magnitude prediction. The system detects whether a topic will grow (0.86 ROC) but cannot "
    "predict how large it will become. The Szabo-Huberman log-linear model (2010) was tested and "
    "achieved R-squared of only 0.22. Random Forest regression, power-law percentile ranges, and "
    "growth multiplier classifiers all failed because at 1-3 posts, 99.5% of word pairs remain "
    "small and the model cannot distinguish a future 10-post topic from a future 22-post topic."
)

para(
    "Exact timing. Regression models for \"days until death\" produced R-squared of -0.5, performing "
    "worse than predicting the mean. Binary classification of the same events works well; regression "
    "on the exact timing does not."
)

para(
    "Revival timing. Predicting when a dead topic will revive achieved only 0.578 ROC AUC. "
    "Revival depends on external real-world events (new developments in an ongoing story), which "
    "cannot be predicted from historical engagement metrics. However, predicting which topics are "
    "the type that revives works at 0.970 ROC AUC."
)

doc.add_heading("Limitations", level=2)
bullet("13 days of data limits rare event learning (approximately 62 topic growth events in test sets)")
bullet("Daily topic granularity; hourly resolution would provide more signal for early detection")
bullet("No natural language processing at the topic level; ongoing stories are distinguished from one-shot events by trajectory shape only")
bullet("Magnitude prediction remains unsolved from early engagement signals")
bullet("All findings are correlational; causal mechanisms cannot be established from observational data")

# ================================================================
# 6. REFLECTION
# ================================================================
doc.add_heading("6. Reflection on AI-Assisted Coding", level=1)

doc.add_heading("Where AI Was Effective", level=2)
para(
    "AI tools enabled rapid prototyping across the entire pipeline. The initial collection "
    "system, history builder, and prediction framework were functional within the first session. "
    "Algorithm selection (Random Forest, VADER, Gini coefficient) and boilerplate code generation "
    "were particularly effective, allowing the developer to focus on experimental design rather "
    "than implementation mechanics."
)

doc.add_heading("Where AI Was Misleading", level=2)

para(
    "Default model assumptions. Both Codex and Claude Code defaulted to Random Forest for every "
    "classification task without suggesting simpler alternatives. Testing revealed that Logistic "
    "Regression outperforms Random Forest on the core emergence detection task, indicating the "
    "signal is linear. This was not discovered until the developer explicitly requested model comparison."
)

para(
    "Redundant feature suggestions. Claude Code recommended using post-level activity state labels "
    "(surging, alive, cooling, dying, dead) as features for topic-level prediction. These labels "
    "are discretised velocity values and showed no separation between growing and non-growing "
    "topics (alive ratio: 0.9x). The developer identified this as redundant information."
)

para(
    "Premature ceiling claims. AI tools declared performance ceilings multiple times during "
    "development. The developer pushed past each one, subsequently discovering the topic lifecycle "
    "pipeline, subreddit spread prediction, death definition analysis, and the ongoing-vs-one-shot "
    "classification, each of which produced strong results."
)

para(
    "Hyperparameter negligence. Default or near-default parameters were used throughout until "
    "the developer requested tuning. Decision Tree performance improved by +0.265 ROC AUC and "
    "Gradient Boosting by +0.122 simply by adjusting tree depth and learning rate."
)

doc.add_heading("Validation Approach", level=2)
bullet("Temporal validation: models trained on days 1-8, tested on days 9-13")
bullet("Feature importance analysis to verify models learn real patterns rather than artefacts")
bullet("Multiple classifier comparison to avoid algorithm selection bias")
bullet("Testing of established research models (Szabo-Huberman) against project data")
bullet("Honest reporting of negative results alongside positive ones")

doc.add_heading("Ethical Considerations", level=2)
bullet("1.1-second rate limiting between all API requests")
bullet("Aggregate analysis only; no individual user identification or profiling")
bullet("Public data only, accessed via Reddit's documented JSON endpoints without authentication")
bullet("Transparent documentation of all AI tools used throughout development")

# ================================================================
# REFERENCES
# ================================================================
doc.add_heading("References", level=1)

refs = [
    "Hutto, C.J. and Gilbert, E.E. (2014) 'VADER: A Parsimonious Rule-based Model for Sentiment Analysis of Social Media Text', Proceedings of the International AAAI Conference on Web and Social Media.",
    "Pedregosa, F. et al. (2011) 'Scikit-learn: Machine Learning in Python', Journal of Machine Learning Research, 12, pp. 2825-2830.",
    "Szabo, G. and Huberman, B.A. (2010) 'Predicting the Popularity of Online Content', Communications of the ACM, 53(8), pp. 80-88.",
]
for r in refs:
    p = doc.add_paragraph(r)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.paragraph_format.left_indent = Cm(1.27)

doc.save(OUTPUT_PATH)
print("Saved: %s" % OUTPUT_PATH)
print("Done!")
